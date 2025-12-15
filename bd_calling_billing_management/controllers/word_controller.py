from odoo.http import Controller, route, request
import json
import logging
import io

_logger = logging.getLogger(__name__)


class DocxPdfConverterController(Controller):

    @route(['/report/download'], type='http', auth="user")
    def report_download(self, data, context=None, token=None):
        """Override report download - convert to DOCX only if is_docx_report=True"""
        try:
            requestcontent = json.loads(data)
            url = requestcontent[0]
            
            # 🚀 FIX: Check if this report should be converted to DOCX
            report_id = self._extract_report_id(url)
            should_convert = self._should_convert_to_docx(report_id)
            
            # Generate normal PDF
            pdf_response = self._normal_download(data, context, token)
            
            # If NOT a DOCX report → return PDF as-is
            if not should_convert:
                _logger.info("Report %s: Returning as PDF (is_docx_report=False)", report_id)
                return pdf_response

            # Convert PDF → DOCX
            _logger.info("Report %s: Converting to DOCX (is_docx_report=True)", report_id)
            pdf_bytes = pdf_response.data
            docx_bytes = self._convert_pdf_to_docx(pdf_bytes)

            # If conversion fails → return PDF
            if not docx_bytes:
                _logger.error("PDF → DOCX conversion failed")
                return pdf_response

            # Build docx filename
            filename = self._extract_filename(pdf_response)
            if filename.endswith(".pdf"):
                filename = filename.replace(".pdf", ".docx")

            return request.make_response(
                docx_bytes,
                headers=[
                    ('Content-Type',
                     'application/vnd.openxmlformats-officedocument.wordprocessingml.document'),
                    ('Content-Disposition', f'attachment; filename="{filename}"')
                ]
            )

        except Exception as e:
            _logger.error("DOCX conversion crash: %s", e)
            return self._normal_download(data, context, token)

    # --------------------------------------------------------------------------
    # 🚀 NEW: Check if report should be converted to DOCX
    def _should_convert_to_docx(self, report_id):
        """Check if ir.actions.report has is_docx_report=True"""
        try:
            if not report_id:
                return False
            
            report = request.env['ir.actions.report'].sudo().browse(int(report_id))
            
            # Check if field exists and is True
            if hasattr(report, 'is_docx_report') and report.is_docx_report:
                return True
            
            return False
        except Exception as e:
            _logger.warning("Could not determine is_docx_report: %s", e)
            return False

    # --------------------------------------------------------------------------
    # 🚀 NEW: Extract report ID from URL
    def _extract_report_id(self, url):
        """Extract report ID from URL like /report/pdf/report_name/123"""
        try:
            # URL format: /report/pdf/module.report_name/record_id
            # or /report/html/module.report_name/record_id
            parts = url.split('/')
            
            if len(parts) < 4:
                return None
            
            report_name = parts[3]  # e.g., "bd_calling_billing_management.without_test_pdf"
            
            # Find report by report_name
            report = request.env['ir.actions.report'].sudo().search([
                ('report_name', '=', report_name)
            ], limit=1)
            
            return report.id if report else None
            
        except Exception as e:
            _logger.warning("Could not extract report ID: %s", e)
            return None

    # --------------------------------------------------------------------------
    # ✅ PDF → DOCX with OPTION 2: Extract totals from bottom-right area
    def _convert_pdf_to_docx(self, pdf_data):
        import io, tempfile, os, fitz, hashlib
        from PIL import Image
        from pdf2docx import Converter
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        try:
            # -----------------------------
            # Step 1: PDF → DOCX with table parsing
            # -----------------------------
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp_pdf, \
                tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp_docx:

                tmp_pdf.write(pdf_data)
                tmp_pdf.flush()

                cv = Converter(tmp_pdf.name)
                cv.convert(
                    tmp_docx.name,
                    start=0,
                    end=None,
                    parse_table=True,           # ✅ Enable to get table data
                    detect_vertical=False,      
                    detect_horizontal=False,    
                    images=False                # images handled manually later
                )
                cv.close()

                docx_bytes = open(tmp_docx.name, "rb").read()

                try:
                    os.unlink(tmp_pdf.name)
                    os.unlink(tmp_docx.name)
                except:
                    pass

            # Load DOCX
            doc = Document(io.BytesIO(docx_bytes))
            pdf = fitz.open(stream=pdf_data, filetype="pdf")
            added_hashes = set()

            # -----------------------------
            # Step 2: Extract totals from bottom-right of each page
            # -----------------------------
            for page_index, page in enumerate(pdf):
                page_width = page.rect.width
                page_height = page.rect.height
                
                # 🚀 OPTION 2: Extract text from bottom-right corner (where totals are)
                # Define the area: right 40% of page, bottom 20% of page
                totals_rect = fitz.Rect(
                    page_width * 0.60,  # Start from 60% right
                    page_height * 0.75,  # Start from 75% down
                    page_width,          # To right edge
                    page_height          # To bottom edge
                )
                
                # Extract text from this area
                totals_text = page.get_text("text", clip=totals_rect).strip()
                
                # If we found totals text, add it to the document
                if totals_text and any(keyword in totals_text.lower() 
                                      for keyword in ['subtotal', 'tax', 'total', 'amount']):
                    _logger.info("Found totals section: %s", totals_text[:100])
                    
                    # Add spacing before totals
                    doc.add_paragraph()
                    
                    # Create a right-aligned paragraph for totals
                    totals_paragraph = doc.add_paragraph()
                    totals_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    
                    # Add the totals text with formatting
                    run = totals_paragraph.add_run(totals_text)
                    run.font.size = Pt(10)
                    run.font.name = 'Arial'
                    
                    # Make "Total:" line bold if present
                    if 'Total:' in totals_text or 'total:' in totals_text:
                        lines = totals_text.split('\n')
                        totals_paragraph.clear()
                        
                        for line in lines:
                            if line.strip():
                                run = totals_paragraph.add_run(line + '\n')
                                run.font.size = Pt(10)
                                run.font.name = 'Arial'
                                
                                # Bold the Total line
                                if 'total:' in line.lower():
                                    run.font.bold = True

            # -----------------------------
            # Step 3: Add images manually (skip watermarks)
            # -----------------------------
            for page_index, page in enumerate(pdf):
                images = page.get_images(full=True)
                page_width = page.rect.width
                page_height = page.rect.height

                for img in images:
                    try:
                        xref = img[0]
                        base = pdf.extract_image(xref)
                        img_bytes = base["image"]
                        width = base.get("width", 0)
                        height = base.get("height", 0)

                        # Skip tiny icons or logos
                        if width < 120 or height < 120:
                            continue

                        # Skip color-indexed (transparent watermark)
                        if base.get("colorspace", "") == "indexed":
                            continue

                        # Skip header/footer watermarks
                        bbox = page.get_image_bbox(xref)
                        if bbox.y0 < page_height * 0.10 or bbox.y1 > page_height * 0.90:
                            continue

                        # Skip duplicates
                        h = hashlib.sha256(img_bytes).hexdigest()
                        if h in added_hashes:
                            continue
                        added_hashes.add(h)

                        img_pil = Image.open(io.BytesIO(img_bytes))
                        if img_pil.mode in ("P", "RGBA"):
                            img_pil = img_pil.convert("RGB")

                        final = io.BytesIO()
                        img_pil.save(final, format="PNG")
                        final.seek(0)

                        max_width = 5.0
                        dpi = 96
                        width_in = min(img_pil.width / dpi, max_width)

                        doc.add_picture(final, width=Inches(width_in))

                    except Exception:
                        continue

                if page_index < len(pdf) - 1:
                    doc.add_page_break()

            pdf.close()

            # -----------------------------
            # Step 4: Save final docx
            # -----------------------------
            output = io.BytesIO()
            doc.save(output)
            return output.getvalue()

        except Exception as e:
            _logger.error("PDF → DOCX conversion failed: %s", e)
            return None

    # --------------------------------------------------------------------------
    def _extract_filename(self, response):
        cd = response.headers.get("Content-Disposition", "")
        if "filename=" in cd:
            return cd.split("filename=")[1].strip('"')
        return "document.docx"

    # --------------------------------------------------------------------------
    def _normal_download(self, data, context, token):
        from odoo.addons.web.controllers.report import ReportController
        return ReportController().report_download(data, context, token)